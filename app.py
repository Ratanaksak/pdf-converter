from flask import Flask, render_template, request, send_file, url_for, redirect
import os
import uuid
import logging
from werkzeug.utils import secure_filename
from pdf2docx import Converter
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer
from docx import Document
import pytesseract
from pdf2image import convert_from_path

# ------------ Config ------------
logging.basicConfig(level=logging.INFO)
app = Flask(__name__)

# Folders (ensure these exist in repo)
BASE_DIR = os.getcwd()
UPLOADS_DIR = os.path.join(BASE_DIR, 'uploads')
STATIC_DIR = os.path.join(BASE_DIR, 'static')
os.makedirs(UPLOADS_DIR, exist_ok=True)
os.makedirs(STATIC_DIR, exist_ok=True)

# Safety & limits
ALLOWED_EXTENSIONS = {'pdf'}
app.config['MAX_CONTENT_LENGTH'] = int(os.environ.get('MAX_CONTENT_LENGTH', 50 * 1024 * 1024))  # default 50MB

# Use env vars so this works on Windows locally and Linux on Render/Docker
TESSERACT_CMD = os.environ.get('TESSERACT_CMD', r"C:\Program Files\Tesseract-OCR\tesseract.exe")
POPPLER_PATH = os.environ.get('POPPLER_PATH', r"C:\poppler-25.07.0\Library\bin")

pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

# ------------ Helpers ------------
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def is_khmer(char):
    return 0x1780 <= ord(char) <= 0x17FF

def contains_khmer(text):
    return any(is_khmer(ch) for ch in text)

def check_khmer_pages(pdf_path):
    khmer_pages = []
    try:
        for page_num, page_layout in enumerate(extract_pages(pdf_path), start=1):
            page_text = ""
            for element in page_layout:
                if isinstance(element, LTTextContainer):
                    page_text += element.get_text()
            if contains_khmer(page_text):
                khmer_pages.append(page_num)
    except Exception as e:
        logging.exception("check_khmer_pages failed: %s", e)
    return khmer_pages

def convert_pdf_to_word(pdf_file_path, dest_docx_path):
    cv = None
    try:
        cv = Converter(pdf_file_path)
        cv.convert(dest_docx_path, start=0, end=None)
    finally:
        if cv:
            cv.close()
    return dest_docx_path

def extract_khmer_text_pages(pdf_file_path, dpi=300):
    ocr_texts = {}
    try:
        pages = convert_from_path(pdf_file_path, dpi=dpi, poppler_path=POPPLER_PATH)
        for i, page in enumerate(pages, start=1):
            text = pytesseract.image_to_string(page, lang="khm")
            if text and text.strip():
                ocr_texts[i] = text.strip()
    except Exception as e:
        logging.exception("extract_khmer_text_pages failed: %s", e)
    return ocr_texts

def paragraph_has_page_break(para):
    # Detect a page break run in the paragraph (best-effort)
    try:
        for run in para.runs:
            xml = run._element.xml
            if 'w:br' in xml and 'type="page"' in xml:
                return True
    except Exception:
        pass
    return False

def is_garbled_text(s):
    if not s or not s.strip():
        return False
    # Replacement char used by some encodings
    if '\ufffd' in s:
        return True
    # If the text contains a high ratio of non-ascii characters but not valid khmer, treat as garbled
    non_ascii = sum(1 for ch in s if ord(ch) > 127)
    ratio = non_ascii / max(len(s), 1)
    # If more than 50% non-ascii but contains no Khmer codepoints => probably garbled
    if ratio > 0.5 and not contains_khmer(s):
        return True
    return False

def replace_khmer_inline_in_docx(docx_path, ocr_texts):
    """
    Best-effort: iterate through paragraphs, detect page breaks and replace paragraphs
    that look garbled with the OCR text for that page.
    """
    try:
        doc = Document(docx_path)
        page_idx = 1
        replaced_pages = set()
        for para in doc.paragraphs:
            # If current page has OCR and this paragraph looks garbled, replace it
            if page_idx in ocr_texts and is_garbled_text(para.text):
                para.text = ocr_texts[page_idx]
                replaced_pages.add(page_idx)

            # increment page index if paragraph contains a page break
            if paragraph_has_page_break(para):
                page_idx += 1

        doc.save(docx_path)
        return replaced_pages
    except Exception as e:
        logging.exception("replace_khmer_inline_in_docx failed: %s", e)
        return set()

def append_ocr_text_at_end(docx_path, ocr_texts):
    try:
        doc = Document(docx_path)
        if ocr_texts:
            doc.add_page_break()
            doc.add_paragraph("⚠️ Khmer text extracted by OCR (best-effort):")
            for pnum in sorted(ocr_texts.keys()):
                doc.add_paragraph(f"--- Page {pnum} ---")
                doc.add_paragraph(ocr_texts[pnum])
            doc.save(docx_path)
    except Exception as e:
        logging.exception("append_ocr_text_at_end failed: %s", e)

# ------------ Routes ------------
@app.route('/')
def index():
    return render_template('pdfconverter.html')

@app.route('/convert', methods=['POST'])
def convert():
    if 'pdfFile' not in request.files:
        return "No PDF file provided.", 400

    uploaded = request.files['pdfFile']
    if uploaded.filename == '' or not allowed_file(uploaded.filename):
        return "Invalid file selected. Please upload a PDF file.", 400

    # Create unique filenames to avoid clashes
    original_name = secure_filename(uploaded.filename)
    uid = uuid.uuid4().hex
    saved_pdf = os.path.join(UPLOADS_DIR, f"{uid}_{original_name}")
    uploaded.save(saved_pdf)
    output_filename = f"{uid}_output.docx"
    output_path = os.path.join(STATIC_DIR, output_filename)

    try:
        # Step 1: detect Khmer pages (fast-ish)
        khmer_pages = check_khmer_pages(saved_pdf)
        logging.info("Detected Khmer pages: %s", khmer_pages)

        # Step 2: convert whole PDF (preserve layout)
        convert_pdf_to_word(saved_pdf, output_path)
        logging.info("pdf2docx conversion saved to %s", output_path)

        # Step 3: if Khmer found -> OCR pages, attempt inline replacement, append OCR text as fallback
        if khmer_pages:
            ocr_texts = extract_khmer_text_pages(saved_pdf, dpi=300)
            logging.info("OCR extracted for pages: %s", list(ocr_texts.keys()))

            replaced_pages = replace_khmer_inline_in_docx(output_path, ocr_texts)
            logging.info("Replaced pages inline: %s", replaced_pages)

            # append remaining OCR text at the end for user reference
            append_ocr_text_at_end(output_path, {p: t for p, t in ocr_texts.items() if p not in replaced_pages})

            message = f"⚠️ PDF contains Khmer characters on pages {khmer_pages}. OCR applied; " \
                      f"inline replacements: {sorted(list(replaced_pages))}. " \
                      f"OCR text appended for other pages. ✅ " \
                      f"<a href='{url_for('download', filename=output_filename)}'>Download Word Document</a>"
            return message
        else:
            return f"✅ Conversion successful! <a href='{url_for('download', filename=output_filename)}'>Download Word Document</a>"

    except Exception as e:
        logging.exception("Conversion failed: %s", e)
        return f"Conversion failed: {str(e)}", 500
    finally:
        # Clean up uploaded PDF (we keep generated DOCX in static)
        try:
            if os.path.exists(saved_pdf):
                os.remove(saved_pdf)
        except Exception:
            pass

@app.route('/download/<filename>')
def download(filename):
    # Basic safety: only serve from static dir and secure the filename
    safe = secure_filename(filename)
    full = os.path.join(STATIC_DIR, safe)
    if not os.path.exists(full):
        return "File not found.", 404
    return send_file(full, as_attachment=True)

# In production, read PORT from environment
if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
